# -*- coding: utf-8 -*-
"""Builds the PM Fellowship strategy + MVP Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN      = RGBColor(0x1D, 0xB9, 0x54)
DARKGREEN  = RGBColor(0x12, 0x6B, 0x36)
GREY       = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x11, 0x11, 0x11)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def _shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexfill)
    tcPr.append(sh)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARKGREEN
    r.font.name = "Calibri"
    # bottom border
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val","single"),("w:sz","8"),("w:space","2"),("w:color","1DB954")):
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12.5); r.font.color.rgb = BLACK
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def kv_bullet(lead, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(lead + " "); r.bold = True; r.font.color.rgb = DARKGREEN
    p.add_run(text)
    return p

def table(headers, rows, widths=None, header_fill="1DB954"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i], header_fill)
        para = hdr[i].paragraphs[0]; run = para.add_run(htext)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]; run = para.add_run(str(val))
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t

# ── Title block ────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Spotify Growth — Breaking the Filter Bubble")
tr.bold = True; tr.font.size = Pt(24); tr.font.color.rgb = DARKGREEN

sub = doc.add_paragraph()
sr = sub.add_run("Problem framing & an AI-native discovery MVP (“Spotify Compass”)")
sr.font.size = Pt(13); sr.font.color.rgb = GREY; sr.italic = True

meta = doc.add_paragraph()
mr = meta.add_run("PM Fellowship Final Project  ·  Growth Team perspective  ·  June 2026")
mr.font.size = Pt(9.5); mr.font.color.rgb = GREY

# ── Executive summary ──────────────────────────────────────────────────────
h1("Executive Summary")
body("Across 27,381 app-store and Reddit reviews, an AI synthesis of that corpus, 27 "
     "primary-research survey responses, and Spotify’s own published research, one "
     "problem dominates: users are trapped in an algorithmic filter bubble. Spotify’s "
     "recommender optimises for past engagement to maximise listening time, which steadily "
     "narrows musical diversity, recycles the same tracks, ignores explicit “hide / don’t "
     "play” feedback, and is contaminated by functional listening (sleep, study). Critically, "
     "users have no control over the “risk level” of recommendations and no way to navigate "
     "by mood.")
body("This document frames the root cause, identifies the target segment (free-tier "
     "listeners, the largest affected group and the direct conversion lever), makes the "
     "business case, and presents a functional AI-native MVP — Spotify Compass — that gives "
     "users an explicit, explainable discovery-controls panel. Because Spotify deprecated its "
     "legacy recommendation and audio-feature APIs in November 2024, the MVP uses a large "
     "language model as the recommendation brain: the very architecture that demonstrates why "
     "AI, not traditional collaborative filtering, is the right tool for this problem.")

# ── 1. Problem framing ─────────────────────────────────────────────────────
h1("1.  Problem Framing")

h2("1.1  Root cause")
body("The visible complaints — “the same songs over and over,” “I can’t influence or reset "
     "them,” “it doesn’t match my energy” — are symptoms. The root cause is a combination of "
     "an optimisation target and an absence of user control:")
kv_bullet("Engagement-maximising objective:",
          "the recommender is tuned to maximise listening time, so it leans on safe, familiar "
          "“bets.” Over time this collapses diversity and produces a self-reinforcing bubble.")
kv_bullet("Ignored negative feedback:",
          "users report that hidden or disliked tracks reappear in other mixes within days.")
kv_bullet("Profile contamination:",
          "functional listening (rain sounds, study loops) bleeds into Discover Weekly and "
          "skews the taste profile.")
kv_bullet("No control surface:",
          "there is no dial for novelty/risk, no taste reset, and no mood-based navigation — "
          "so the only escape is manually hunting for new music, which is high-friction.")
body("In short: the system decides for the user and cannot be steered. That is a product-design "
     "gap, not merely a model-quality gap.", italic=True)

h2("1.2  The evidence is consistent across every source")
table(
    ["Source", "What it shows"],
    [
        ["Review corpus (27,381)", "‘Recommendation quality’, ‘shuffle/repetition’ and "
         "‘missing features’ are the top complaint themes; 3,394 reviews voice unmet needs."],
        ["AI synthesis", "Most-affected segment = free-tier (12.5% negative); highest negative "
         "rate = programmed listeners (12.7%); churned users (12.3%) cite recommendations + price."],
        ["Primary survey (27)", "High ‘bubble’ scores; recurring ‘can’t influence or reset,’ "
         "‘ignores my mood’; two users independently asked for an ‘AI DJ that understands my mood.’"],
        ["Spotify research / brief", "Confirms the engagement-driven filter bubble, ignored "
         "negative feedback, contamination, and the lack of ‘risk level’ controls."],
    ],
    widths=[1.6, 5.4],
)

h2("1.3  Target segment (lead: free-tier listeners)")
body("We lead with free-tier listeners and treat at-risk/churning users as the secondary "
     "beneficiary:")
kv_bullet("Free-tier (primary):",
          "the largest real-volume affected group (12.5% negative) and the direct revenue "
          "lever. The discovery bubble compounds with ad friction to suppress perceived value "
          "and push these users to YouTube Music (“more control, cheaper”).")
kv_bullet("Programmed listeners (secondary):",
          "highest negative rate (12.7%); they live inside autoplay/radio, exactly where the "
          "bubble is worst.")
kv_bullet("Previously active (secondary):",
          "12.3% negative and actively switching to competitors, citing recommendations + price.")

h2("1.4  Why solving it makes business sense")
body("Discovery dissatisfaction is a Growth problem, not just a satisfaction problem:")
bullet("Better, steerable discovery raises engagement (sessions, completion, saves), which is "
       "the leading indicator of free→Premium conversion.", bold_lead="Conversion: ")
bullet("Competitors are winning on “control.” Closing that gap reduces churn and supports "
       "win-back of previously-active users.", bold_lead="Retention: ")
bullet("Advanced controls (unlimited resets, sandbox, save-as-playlist, higher novelty depth) "
       "are a natural Premium upsell — a conversion hook built into the feature itself.",
       bold_lead="Monetisation: ")
body("Sizing is deliberately directional rather than a fabricated precise figure: on a free "
     "base of hundreds of millions of monthly actives, even single-digit-basis-point lifts in "
     "free→Premium conversion translate into material ARR. Primary metrics: free→Premium "
     "conversion rate and D30 retention; secondary: sessions/week, save rate, skip rate, churn, "
     "and competitive-switch intent in NPS detractor reasons.")

# ── 2. The MVP ─────────────────────────────────────────────────────────────
h1("2.  The Solution — “Spotify Compass” MVP")

h2("2.1  Concept")
body("Spotify Compass is an embedded discovery-controls panel layered onto the listening "
     "experience. Instead of a black box that decides for you, it hands the user a small set of "
     "explicit, legible controls and explains every recommendation. It is a functional "
     "prototype: it authenticates with real Spotify accounts, returns real catalogue tracks, "
     "and saves real playlists.")

h2("2.2  The discovery-controls panel")
kv_bullet("Mood selector:", "choose the emotional/energy target directly (calm, focus, "
          "workout-hype, melancholy…).")
kv_bullet("Novelty / risk dial (1–5):", "from “stay in my comfort zone” to “break my bubble” — "
          "novelty as a first-class control, not a side effect.")
kv_bullet("Sandbox / reset taste:", "ignore listening history for this session; nothing here "
          "contaminates the long-term profile.")
kv_bullet("Block artist (sticky):", "negative feedback that is honoured immediately and "
          "persists across regenerations.")
kv_bullet("Per-track “why”:", "each pick carries a one-line reason, restoring transparency "
          "and trust.")

h2("2.3  How the controls map to the unmet needs")
table(
    ["Control", "Unmet need it solves"],
    [
        ["Novelty / risk dial", "Control over recommendation “risk level”"],
        ["Sandbox / reset taste", "Reset taste profile + stop playlist contamination"],
        ["Mood selector", "Mood-based direct navigation"],
        ["Block artist (sticky)", "Honour negative feedback"],
        ["Per-track “why”", "Transparency / rebuild trust"],
    ],
    widths=[2.4, 4.6],
)

h2("2.4  Architecture")
body("The flow turns the deprecation of Spotify’s legacy recommender into the design itself:")
bullet("User authenticates via Spotify OAuth (real account, free or Premium).", bold_lead="1. ")
bullet("User sets the controls (mood, novelty, context, sandbox, blocked artists).", bold_lead="2. ")
bullet("An LLM (deepseek-v4-flash via the OpenCode API) takes those controls plus an optional "
       "real taste snapshot and returns a structured, explained track list — with novelty as a "
       "parameter and blocked artists guaranteed absent.", bold_lead="3. ")
bullet("Each suggestion is resolved to a real track via the Spotify Search API.", bold_lead="4. ")
bullet("The queue renders with album art and rationale; one click saves it as a real private "
       "Spotify playlist.", bold_lead="5. ")

# ── 3. Why AI ──────────────────────────────────────────────────────────────
h1("3.  Why AI Is Uniquely Suited")

h2("3.1  Why traditional recommendation systems are insufficient")
body("Collaborative filtering and matrix-factorisation recommenders are powerful but "
     "structurally mismatched to this problem:")
bullet("They optimise a single objective (engagement), so they converge on popularity and "
       "familiarity — the bubble is a feature of the maths, not a bug.")
bullet("They are opaque: no explanation, so users distrust the output (“based on who paid for "
       "placement”).")
bullet("They accept only coarse feedback (like/skip) and cannot parse natural-language intent "
       "or real-time context such as mood or activity.")
bullet("They offer no novelty control and no taste reset.")
body("Tellingly, Spotify deprecated its /recommendations, audio-features and related-artists "
     "endpoints for new apps in November 2024 — the legacy primitives are themselves no longer "
     "the path forward.", italic=True)

h2("3.2  What AI unlocks that was previously difficult")
kv_bullet("Natural-language intent + mood:", "the model interprets “winding down, something "
          "calm but new to me” directly.")
kv_bullet("Novelty as a tunable parameter:", "exploration becomes a dial the user owns, not an "
          "emergent side effect.")
kv_bullet("Explainability:", "a specific reason per track, which rebuilds trust.")
kv_bullet("Sticky negative feedback:", "blocked artists are excluded deterministically and "
          "instantly.")
kv_bullet("Sandboxed taste reset:", "explore a new direction without polluting the long-term "
          "profile — previously hard because the profile and the recommender were one and the "
          "same.")

h2("3.3  How AI changes the user experience")
body("The experience shifts from passive, opaque, and uncontrollable lean-back consumption to "
     "an active, transparent, user-steered co-pilot. The user moves from “the algorithm decides "
     "and I can’t change it” to “I set the intent and the risk, and I can see why.”")

h2("3.4  Traditional recsys vs. the AI-native approach")
table(
    ["Dimension", "Traditional recsys", "AI-native (Compass)"],
    [
        ["Objective", "Maximise engagement", "Serve the user’s stated intent"],
        ["Novelty", "Emergent side effect", "Explicit, tunable dial"],
        ["Input", "Clicks, skips", "Natural language + mood + context"],
        ["Feedback", "Coarse, often ignored", "Honoured instantly & persistently"],
        ["Transparency", "Opaque", "Per-track explanation"],
        ["Taste reset", "Not possible", "Sandbox mode"],
    ],
    widths=[1.5, 2.6, 2.9],
)

# ── 4. Metrics & rollout ───────────────────────────────────────────────────
h1("4.  Success Metrics & Rollout")
kv_bullet("North-star:", "free→Premium conversion rate; D30 retention.")
kv_bullet("Engagement proxies:", "sessions/week, listening completion, save rate, skip-rate "
          "reduction.")
kv_bullet("Discovery health:", "share of distinct new artists per user per week; reduction in "
          "“repetition” complaint volume.")
kv_bullet("Rollout:", "ship to a free-tier cohort as an opt-in panel → A/B against control on "
          "engagement & conversion → gate the advanced controls (depth, unlimited resets, "
          "save-as-playlist) behind Premium as the conversion hook.")

# ── Appendix ───────────────────────────────────────────────────────────────
h1("Appendix — Representative User Voices")
for q in [
    "“I hide songs but they keep reappearing in other mixes — I can’t escape the loop.”",
    "“Idk how to refresh my recommendation to get new music.”",
    "“Spotify don’t match my energy, ig that’s too much to ask for.”",
    "“I’d love a smarter AI DJ that actually understands my mood and creates playlists in real time.”",
    "“YouTube Music gives me more control and is cheaper.”",
]:
    p = doc.add_paragraph(); r = p.add_run(q); r.italic = True; r.font.color.rgb = GREY
    p.paragraph_format.left_indent = Inches(0.3)

out = "mvp/Spotify_Growth_Discovery_MVP.docx"
doc.save(out)
print("saved", out)
